"""DOCX document parser.

Ported from 김보윤's RFP_RAG_project extractors.py.
Uses python-docx to extract text and tables from DOCX files.
"""
import os
from typing import List

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from bidflow.domain.models import ParsedChunk, ParsedTable


class DOCXParser:
    """DOCX 텍스트/테이블 추출 파서"""

    def parse(
        self, file_path: str, chunk_size: int = 500, chunk_overlap: int = 50
    ) -> List[ParsedChunk]:
        text_content = self._extract_text(file_path)
        if not text_content:
            print("[DOCXParser] No text extracted")
            return []

        from bidflow.parsing.preprocessor import TextPreprocessor
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        preprocessor = TextPreprocessor()
        normalized = preprocessor.normalize(text_content)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""],
        )
        split_texts = splitter.split_text(normalized)

        chunks = []
        for i, text in enumerate(split_texts):
            chunks.append(
                ParsedChunk(
                    chunk_id=f"docx_chunk_{i}",
                    text=text,
                    page_no=1,
                    metadata={
                        "source": "docx",
                        "filename": os.path.basename(file_path),
                        "file_path": file_path,
                        "chunk_index": i,
                    },
                )
            )
        return chunks

    def extract_tables(self, file_path: str) -> List[ParsedTable]:
        tables = []
        try:
            doc = DocxDocument(file_path)
            for j, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    tables.append(
                        ParsedTable(
                            table_id=f"docx_t{j + 1}",
                            page_no=1,
                            rows=rows,
                            metadata={
                                "source": "docx",
                                "filename": os.path.basename(file_path),
                            },
                        )
                    )
        except Exception as e:
            print(f"[DOCXParser] Table extraction failed: {e}")
        return tables

    def _extract_text(self, file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            parts: list[str] = []

            for element in doc.element.body:
                tag = element.tag

                if tag.endswith("}p") or tag == "p":
                    text = element.text
                    if not text:
                        text = "".join(
                            t.text for t in element.iter(qn("w:t")) if t.text
                        )
                    if text and text.strip():
                        parts.append(text.strip())

                elif tag.endswith("}tbl") or tag == "tbl":
                    rows: list[list[str]] = []
                    for tr in element.iter(qn("w:tr")):
                        cells: list[str] = []
                        for tc in tr.iter(qn("w:tc")):
                            cell_text = "".join(
                                t.text for t in tc.iter(qn("w:t")) if t.text
                            )
                            cells.append(cell_text.strip())
                        if cells:
                            rows.append(cells)

                    if rows:
                        # col_path 스타일 직렬화
                        headers = rows[0] if rows else []
                        for row in rows[1:]:
                            row_parts = []
                            for col_idx, cell in enumerate(row):
                                header = headers[col_idx] if col_idx < len(headers) else f"col_{col_idx}"
                                if cell.strip():
                                    row_parts.append(f"{header}: {cell}")
                            if row_parts:
                                parts.append(" | ".join(row_parts))

            return "\n".join(parts)
        except Exception as e:
            print(f"[DOCXParser] Text extraction failed: {e}")
            return ""
